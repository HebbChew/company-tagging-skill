#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""分析报告 docx 版式构建器（公文规范）。改内容改 CONTENT，改版式改样式函数。"""
import os
import docx, re
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x6F, 0xB2)
GRAY = RGBColor(0x59, 0x59, 0x59)

def set_run(r, cn='仿宋', en='Times New Roman', size=12, bold=False, color=None):
    r.font.name = en
    r._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color: r.font.color.rgb = color

def para(doc, text='', style_cn='仿宋', size=12, bold=False, align=None, indent=True,
         space_after=6, line=1.5, color=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    for seg in re.split(r'(\*\*.+?\*\*)', text):
        if not seg: continue
        if seg.startswith('**'):
            set_run(p.add_run(seg[2:-2]), style_cn, size=size, bold=True, color=color)
        else:
            set_run(p.add_run(seg), style_cn, size=size, bold=bold, color=color)
    return p

def h1(doc, text):
    p = para(doc, text, style_cn='黑体', size=15, bold=True, indent=False, space_after=8)
    p.paragraph_format.space_before = Pt(14)
    return p

def h2(doc, text):
    p = para(doc, text, style_cn='黑体', size=13, bold=True, indent=False, space_after=6)
    p.paragraph_format.space_before = Pt(8)
    return p

def caption(doc, text):
    return para(doc, text, style_cn='楷体', size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, color=GRAY)

def add_table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(header):
        cell = hdr.cells[i]
        cell._tc.get_or_add_tcPr().append(parse_shading())
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(str(h)), '黑体', size=10.5, bold=True, color=RGBColor(255, 255, 255))
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            set_run(p.add_run(str(v)), '宋体', size=10.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    return t

from docx.oxml import parse_xml
def parse_shading():
    return parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="1F6FB2"/>')

def add_chart(doc, path, cap, width=6.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    caption(doc, cap)

doc = docx.Document()
# 页面：A4，公文页边距
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(3.0), Cm(2.5)
sec.left_margin, sec.right_margin = Cm(2.8), Cm(2.6)

# ===== 封面块 =====
for _ in range(4): doc.add_paragraph()
para(doc, '闵行区生物医药企业全景标注分析报告', '黑体', 24, True, WD_ALIGN_PARAGRAPH.CENTER, indent=False)
doc.add_paragraph()
para(doc, '—— P1 强相关企业 2698 家 · 领域标签 / 一句话简介 / 母公司 / 相关度', '楷体', 13, False, WD_ALIGN_PARAGRAPH.CENTER, indent=False)
for _ in range(6): doc.add_paragraph()
para(doc, '数据截至：2026 年 8 月 27 日', '仿宋', 13, False, WD_ALIGN_PARAGRAPH.CENTER, indent=False)
para(doc, '方法：模型知识初判 + 规则打标 + 梯度搜索核验（1829 家完成核验）', '仿宋', 12, False, WD_ALIGN_PARAGRAPH.CENTER, indent=False)
doc.add_page_break()

# ===== 摘要 =====
h1(doc, '摘  要')
para(doc, '本次对名单 2698 家企业完成领域标签、一句话简介、母公司识别、相关度分级四项标注。核心数字如下：')
add_table(doc,
    ['指标', '家数', '占比/说明'],
    [['生物医药核心主业（R0）', '1406', '52%'],
     ['泛健康（R1）', '111', '有健康监管资质的消费健康/合成生物非医药应用'],
     ['供应链与产业配套（R2）', '164', '含医药供应链上游配套（阀门/气体/装备部件）'],
     ['大健康边缘（R3）', '114', '无资质美妆/食品营销、养生等，建议降级观察'],
     ['完全无关（R4）', '28', '建议移出名单'],
     ['完成搜索核验', '1829', '68%；其中 892 家获得带证据的行业地位表述'],
     ['识别出母公司', '481', '36 个集团在闵行多主体布局'],
     ['无公开信息', '671', '其中 ≥1000 万高资本段经多角度复找后仅 31 家确证为壳']],
    widths=[5.2, 2.2, 8.6])
doc.add_paragraph()
para(doc, '一句话结论：名单"头部扎实、腰部可用、尾部需瘦身"——核心企业与前沿赛道标得清、靠得住；边界企业与小微壳企业已单独成清单，供名单管理决策。', indent=True)
doc.add_page_break()

# ===== 一、名单总体成色 =====
h1(doc, '一、名单总体成色')
para(doc, '按"与生物医药产业的相关度"五级分类（R0 核心至 R4 无关），判级均有证据或明确口径（泛健康与无关的分界为是否持有药/械/保健食品等健康监管资质）。')
add_chart(doc, 'full/charts/图2_相关度结构.png', '图 1  相关度结构分布（已判定 1733 家）')
para(doc, '五级分类的设计意图：R0 是名单主体；R1 保留标注（泛健康有价值但不计入医药核心）；R2 单列（供应链上游与平台，招商口径单独统计）；R3 降级观察；R4 建议移出。')
add_table(doc, ['相关度', '家数', '含义', '名单建议'],
    [['R0 核心', '1406', '药/械/IVD/CXO/流通/医疗服务', '名单主体'],
     ['R1 泛健康', '111', '有健康资质的消费品/合成生物非医药应用', '保留标注'],
     ['R2 配套', '164', '供应链上游+产业服务平台', '单列统计'],
     ['R3 边缘', '114', '无资质美妆/食品营销、养生', '降级观察'],
     ['R4 无关', '28', '纺织检测/宠物殡葬/个护电器等', '建议移出']],
    widths=[2.4, 1.8, 7.0, 4.0])

# ===== 二、产业结构 =====
doc.add_page_break()
h1(doc, '二、产业结构：器械与流通为底盘，前沿赛道密集冒头')
add_chart(doc, 'full/charts/图1_一级领域分布.png', '图 2  一级领域分布（命中标签的 1054 家）')
para(doc, '医疗器械（377 家）与医药流通供应链（251 家）构成名单底盘；药物（153 家）与 CXO（104 家）为中坚；合成生物、脑机接口、核酸药物等前沿方向在小微段密集出现。')
h2(doc, '（一）细分赛道 TOP10 与代表企业')
add_table(doc, ['细分赛道', '家数', '代表企业（地位均经核验）'],
    [['医药流通与供应链', '245', '生生医药冷链（温控供应链全国第一）、朗脉洁净（国家级小巨人）'],
     ['医用材料与耗材', '100', '士卓曼（全球牙科种植体龙头）、上药第一生化'],
     ['IVD-平台技术', '77', '思路迪（全球首款外泌体卵巢癌 IVD）、之江生物（688317）'],
     ['康复与家用医疗', '53', '互邦智能（轮椅国标起草单位）'],
     ['合成生物学', '40', '弈柯莱（国投聚力 4.5 亿战投）、凯赛（688065）'],
     ['小分子药物', '38', '汇伦医药（2026 递表港交所）'],
     ['细胞与基因治疗', '37', '华夏源、健信生物'],
     ['IVD-临床应用', '33', '捷易生物（遗传病诊断全链条）'],
     ['介入与植入', '27', '捍宇医疗（国产首款）、励楷科技（D 轮近 5 亿）'],
     ['手术系统与机器人', '19', '术锐（国内首款单孔）、朗合医疗（支气管镜机器人）']],
    widths=[4.0, 1.6, 9.6])

# ===== 三、前沿赛道亮点 =====
doc.add_page_break()
h1(doc, '三、前沿赛道亮点企业（本次核验新挖出）')
add_table(doc, ['方向', '亮点企业'],
    [['脑机接口', '太易生命科学（美国 TAE 旗下 BNCT）、玮脑智联（心玮 06609 全资）、明视医疗（明视脑机系视觉重建，天使轮 1.5 亿）、曦涟医疗（高瓴/蓝驰背景）、东瑙医疗（海河实验室基金独投）'],
     ['CGT/核酸', '锋寻生物（体内 CAR-T，交大蔡宇伽系近 6000 万种子轮）、艾码申华（外泌体 siRNA，FDA 孤儿药）、靖因药业（二次递表港交所）、恒诺康（吉利德/辉瑞团队）'],
     ['AI 制药', '生鹜医药（天鹜系）、数因信科（嘉必优联营）、海健百世（复旦马剑鹏团队）、迪普深合（交大点击化学+深势科技）'],
     ['器械创新', '心瓴医工（离体心脏温血灌流）、匠鑫医疗（国产首款电动旋磨）、影为医疗（上海电气/天智航战投）、慧威医疗（腔内影像机器人三类证）'],
     ['供应链配套', '珐成制药（至纯科技全资）、东富龙生物试剂（培养基 FDA DMF）、仁科生物（10x Genomics 中国区总代）']],
    widths=[2.8, 12.4])
para(doc, '上述企业的地位表述全部带证据链接（见终稿"证据摘要"列），可直接追溯。', indent=True)

# ===== 四、集团化布局 =====
h1(doc, '四、集团化布局：481 家有母公司，36 个集团多主体落子闵行')
add_table(doc, ['类型', '代表'],
    [['本土龙头全链条', '东富龙系 7 家（装备/试剂/包装/检测/工程/耗材全覆盖）'],
     ['上市药企在沪抓手', '信达生物（研发中心+零售）、之江系 3 家、圣湘系 3 家、苑东系 3 家、先声系、恒瑞（核药平台 2026 新设）、乐普系'],
     ['外资在华主体', 'Eurofins 系 4 家、士卓曼×2、强生/科赴系、美敦力/柯惠系、3M（已分拆 Solventum）、圣戈班、Greiner、义获嘉、InBody'],
     ['跨界集团', '好丽友（桉莱，Orion 77%）、康桥资本系（汇衡功道 mRNA 基地）、宝洁/吉列（博朗，R4）']],
    widths=[4.0, 11.2])

# ===== 五、资本与成色 =====
doc.add_page_break()
h1(doc, '五、资本与成色：1000 万是分水岭')
add_chart(doc, 'full/charts/图3_资本档地位率.png', '图 3  注册资本档 × 有行业地位企业占比（已搜索企业）')
add_table(doc, ['注册资本档', '家数', '已搜索', '有地位词', '有母公司', '无信息'],
    [['≥5000 万', '152', '100%', '89%', '57%', '1%'],
     ['2000-5000 万', '141', '100%', '82%', '48%', '6%'],
     ['1000-2000 万', '290', '100%', '84%', '41%', '7%'],
     ['500-1000 万', '389', '100%', '39%', '21%', '49%'],
     ['100-500 万', '1164', '60%', '40%*', '—', '—'],
     ['<100 万', '562', '28%', '35%*', '—', '—']],
    widths=[3.2, 2.0, 2.4, 2.6, 2.6, 2.2])
para(doc, '*按已搜索企业计。1000 万以上企业 84% 以上有可查证的行业地位；500 万以下无信息率陡增。这直接支撑"≤100 万不搜"的成本决策——搜了也大概率是空白。', indent=True)
para(doc, '本次对 127 家"高资本却无信息"企业做了多角度复找（信用代码/地址/法人/长尾信源），96 家（75%）挖到实质信息，找回智孔生物（先声/晨壹 B 轮系）、正大系两家新设研发主体、奎瑞尔（TIL 专利）、捷鹿（生物传感）等；剩余 31 家可确证为壳。', indent=True)

# ===== 六、名单质量与时效 =====
h1(doc, '六、名单质量与时效警示')
para(doc, '**（一）边界企业 142 家**（清单单独成表）：R4 建议移出 28 家（纺织品检测×2、宠物殡葬、水族电商、家庭园艺肥料、指纹考勤机等）；R3 降级观察 114 家（美妆营销 62、食品营销 36、养生 2、其他 13）。')
para(doc, '**（二）无公开信息 671 家**：特征为整百万注册资本聚集（462 家恰好 100 万）、集中登记地、2025-2026 批量新设。建议与招商口径对齐决定去留。')
para(doc, '**（三）时效事项**：达颀医疗（W&H 在华）报道拟解散清算；中美施贵宝股权易主高瓴系联合赛尔；因美纳母公司列入不可靠实体清单（2025.2）；诺唯赞上海主体更名嘉有健。引用前请按终稿备注复核时点。')

# ===== 七、方法与可信度 =====
h1(doc, '七、方法与可信度（为什么这份数据敢用）')
para(doc, '三层标注体系：模型知识初判（试点实测：地位词 92% 准确、母公司"确定"档 94% 准确）→ 规则打标（零成本、词表校验归一）→ 梯度搜索核验（搜索深度跟着"模型的无知程度"走，知名企业只做 1 次时效核验，陌生但重要的企业深挖 2-3 次）。')
para(doc, '四条红线：无信息企业一个字不猜（671 家空白是有意为之）；地位词必须带权威证据链接（892 家全带）；身份同一性 2/3 规则（全程零张冠李戴）；R3/R4 不打领域标签。')
para(doc, '已知残留风险（如实披露）：「登记」档流通标签约三分之二准确（工商登记口径未经实证）；196 家存疑待核母公司建议天眼查穿透复核；本次核验基于搜索摘要级证据，不等于全文尽调。')

# ===== 八、建议 =====
h1(doc, '八、建议动作')
para(doc, '（一）名单管理：R4 移出、R3 降级观察、671 家无信息企业按招商口径决定去留。')
para(doc, '（二）招商跟进：第三节亮点企业（脑机/CGT/AI 制药/器械创新）建议纳入重点项目跟踪表。')
para(doc, '（三）数据资产：196 家存疑母公司走天眼查批量穿透；每季度对已核验企业跑一次时效核验（分拆/更名/获批）；无信息清单作为招商摸排任务来源。')
doc.add_paragraph()
para(doc, '附件：终稿 Excel（原表+标注 / 终稿 双 sheet）｜质检清单 3 份｜方法文档与全程缓存（22.企业打标/）', '楷体', 10.5, False, indent=False, color=GRAY)

doc.save(os.environ.get('MHB_REPORT_DOCX', '分析报告_企业标注.docx'))
print('设计版报告已生成')
